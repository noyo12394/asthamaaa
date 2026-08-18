from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "manuscript"
ASSET_DIR = OUTPUT_DIR / "assets"
DOCX_PATH = OUTPUT_DIR / "PASS_Terrain_Smoke_Manuscript_Draft.docx"
SCREENSHOT_PATH = ASSET_DIR / "terrain-smoke-lab.png"
RESULTS_FIGURE_PATH = ASSET_DIR / "allentown-worked-example.png"

BLACK = "000000"
MUTED = "555555"
BORDER = "DADCE0"


def set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(BLACK)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single" if edge in {"top", "bottom", "insideH"} else "nil")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), BORDER)
    inside_v = borders.find(qn("w:insideV"))
    if inside_v is None:
        inside_v = OxmlElement("w:insideV")
        borders.append(inside_v)
    inside_v.set(qn("w:val"), "nil")


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
    set_table_borders(table)


def add_numbering_definition(document: Document, *, bullet: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\u25cf" if bullet else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "276")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(document: Document, text: str, num_id: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    paragraph._p.get_or_add_pPr().append(num_pr)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_body(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        lead.font.name = "Arial"
        lead.font.size = Pt(11)
        rest = paragraph.add_run(text[len(bold_lead):])
        rest.font.name = "Arial"
        rest.font.size = Pt(11)
    else:
        run = paragraph.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_with_next = False
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(paragraph)


def add_table(
    document: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    widths_dxa: Sequence[int],
    numeric_columns: set[int] | None = None,
) -> None:
    numeric_columns = numeric_columns or set()
    table = document.add_table(rows=1, cols=len(headers))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER if index in numeric_columns else WD_ALIGN_PARAGRAPH.LEFT)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if index in numeric_columns else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[index], str(value), align=align)
    set_table_geometry(table, widths_dxa)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def create_results_figure() -> None:
    width, height = 2100, 660
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
    bold_path = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")

    def font(size: int, bold: bool = False):
        path = bold_path if bold else font_path
        return ImageFont.truetype(str(path), size) if path.exists() else ImageFont.load_default()

    def centered_text(box: tuple[int, int, int, int], text: str, size: int, *, bold: bool = False, fill="#111827"):
        fnt = font(size, bold)
        bounds = draw.textbbox((0, 0), text, font=fnt)
        x = box[0] + (box[2] - box[0] - (bounds[2] - bounds[0])) / 2
        y = box[1] + (box[3] - box[1] - (bounds[3] - bounds[1])) / 2
        draw.text((x, y), text, font=fnt, fill=fill)

    centered_text((0, 5, width, 70), "Allentown worked example: a visible contrast that did not generalize", 42, bold=True)
    panel_width = 650
    panels = [(35, 90, 685, 625), (725, 90, 1375, 625), (1415, 90, 2065, 625)]
    for panel in panels:
        draw.rounded_rectangle(panel, radius=14, outline="#D1D5DB", width=2, fill="#FFFFFF")

    centered_text((35, 105, 685, 160), "Current terrain contrast", 32, bold=True)
    centered_text((35, 155, 685, 200), "OED = +5.2 ug/m3", 27, fill="#374151")
    baseline_y = 515
    chart_top = 220
    chart_height = 260
    for x, label, value, color in ((210, "Lowland", 8.8, "#F59E0B"), (465, "Highland", 3.6, "#2563EB")):
        bar_height = int((value / 10) * chart_height)
        draw.rectangle((x - 65, baseline_y - bar_height, x + 65, baseline_y), fill=color)
        centered_text((x - 80, baseline_y - bar_height - 45, x + 80, baseline_y - bar_height), f"{value:.1f}", 28, bold=True)
        centered_text((x - 110, baseline_y + 5, x + 110, baseline_y + 55), label, 25)
    draw.line((110, baseline_y, 610, baseline_y), fill="#6B7280", width=2)
    draw.text((55, 260), "PM2.5", font=font(23), fill="#4B5563")

    centered_text((725, 105, 1375, 160), "Temporal holdout ablation", 32, bold=True)
    centered_text((725, 155, 1375, 200), "Terrain lift = -8.2%", 27, fill="#9B1C1C")
    for x, label, value, color in ((900, "Baseline", 7.93, "#6B7280"), (1160, "+ terrain", 8.58, "#0F766E")):
        bar_height = int((value / 10) * chart_height)
        draw.rectangle((x - 65, baseline_y - bar_height, x + 65, baseline_y), fill=color)
        centered_text((x - 80, baseline_y - bar_height - 45, x + 80, baseline_y - bar_height), f"{value:.2f}", 28, bold=True)
        centered_text((x - 110, baseline_y + 5, x + 110, baseline_y + 55), label, 25)
    draw.line((800, baseline_y, 1300, baseline_y), fill="#6B7280", width=2)
    draw.text((745, 260), "RMSE", font=font(23), fill="#4B5563")

    centered_text((1415, 105, 2065, 160), "Evidence pipeline", 32, bold=True)
    stages = ["29-cell terrain grid", "Smoke + PM2.5 + weather", "OED + Spearman rho", "Held-out RF ablation"]
    for index, label in enumerate(stages):
        top = 190 + index * 98
        draw.rounded_rectangle((1515, top, 1965, top + 62), radius=10, outline="#9CA3AF", width=2, fill="#F9FAFB")
        centered_text((1515, top, 1965, top + 62), label, 24)
        if index < len(stages) - 1:
            draw.line((1740, top + 62, 1740, top + 88), fill="#4B5563", width=3)
            draw.polygon([(1732, top + 82), (1748, top + 82), (1740, top + 94)], fill="#4B5563")

    image.save(RESULTS_FIGURE_PATH, quality=95)


def configure_document(document: Document) -> tuple[int, int]:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.widow_control = True

    settings = {
        "Heading 1": (20, BLACK, 20, 6),
        "Heading 2": (16, BLACK, 18, 6),
        "Heading 3": (14, "434343", 16, 4),
    }
    for style_name, (size, color, before, after) in settings.items():
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    bullet_num_id = add_numbering_definition(document, bullet=True)
    decimal_num_id = add_numbering_definition(document, bullet=False)
    return bullet_num_id, decimal_num_id


def build_document() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    if not SCREENSHOT_PATH.exists():
        raise FileNotFoundError(f"Missing application screenshot: {SCREENSHOT_PATH}")
    create_results_figure()

    document = Document()
    bullet_num_id, decimal_num_id = configure_document(document)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("PASS Terrain-Aware Smoke Atlas")
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.bold = False
    run.font.color.rgb = RGBColor.from_string(BLACK)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.keep_with_next = True
    subtitle_run = subtitle.add_run("An uncertainty-first method for testing orographic PM2.5 contrasts in a public environmental-health interface")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(15)
    subtitle_run.font.color.rgb = RGBColor.from_string("434343")

    for line in (
        "Noyonica Chatterjee",
        "Lehigh University",
        "Correspondence: noc224@lehigh.edu",
        "Research manuscript draft | August 2026",
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        if line.startswith("Research manuscript"):
            run.font.color.rgb = RGBColor.from_string(MUTED)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(14)
    note.paragraph_format.space_after = Pt(0)
    note_run = note.add_run(
        "Draft status. This manuscript defines the research contribution and reports a worked application example. It does not yet claim a completed multi-city epidemiologic validation or clinical decision support system."
    )
    note_run.font.name = "Arial"
    note_run.font.size = Pt(10)
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor.from_string(MUTED)

    document.add_page_break()

    add_heading(document, "Abstract", 1)
    add_body(
        document,
        "Background. Public air-quality maps usually display regional pollution surfaces, while complex-terrain studies show that elevation, valley geometry, inversions, and smoke transport can produce neighborhood-scale contrasts. The practical gap is not a lack of maps or machine-learning models. It is the absence of an accessible workflow that lets users inspect whether terrain is actually informative at a selected place, while preserving provenance and negative results."
    )
    add_body(
        document,
        "Objective. We developed PASS Terrain-Aware Smoke Atlas, an interactive research workbench that links gridded elevation, modeled PM2.5, meteorology, and NOAA Hazard Mapping System smoke polygons. The central contribution is a falsifiable local analysis rather than a universal terrain correction."
    )
    add_body(
        document,
        "Methods. For a user-selected location, the system samples a 29-cell circular grid, classifies cells into lowland, transition, and highland thirds, and computes the Orographic Exposure Differential (OED): median lowland PM2.5 minus median highland PM2.5. It also estimates the Spearman association between elevation and PM2.5 and stratifies hourly contrasts by observed smoke overhead. A random-forest ablation compares a meteorology-and-smoke baseline with a model that additionally includes elevation, topographic position, and ruggedness. Models are trained on earlier six-hour blocks and evaluated on the final 24 UTC hours, which are never used for fitting."
    )
    add_body(
        document,
        "Worked example. In an August 2026 Allentown, Pennsylvania query, the current OED was +5.2 micrograms per cubic meter and the elevation-PM2.5 Spearman coefficient was -0.43. However, terrain worsened held-out RMSE from 7.93 to 8.58, a predictive lift of -8.2%. The interface retained and explained this null result instead of converting a single spatial snapshot into a generalized claim."
    )
    add_body(
        document,
        "Conclusion. The proposed novelty is a transparent, uncertainty-first bridge between terrain-aware environmental modeling and public-facing exploration. The workbench operationalizes a new interpretable contrast metric, tests terrain through explicit ablation, and makes source status and model failure visible. A preregistered multi-city, monitor-anchored evaluation is required before health or policy inference."
    )
    add_body(document, "Keywords: PM2.5; wildfire smoke; complex terrain; air quality; asthma; random forest; uncertainty; geospatial visualization")

    add_heading(document, "1. Introduction", 1)
    add_body(
        document,
        "Fine particulate matter with aerodynamic diameter at or below 2.5 micrometers (PM2.5) is a central environmental-health concern because particles penetrate deeply into the respiratory tract and wildfire smoke episodes can sharply increase short-term exposure. Reviews consistently associate wildfire smoke with worsening asthma outcomes, and recent interrupted time-series evidence from Ontario reported a 23.6% increase in asthma-related emergency visits one day after severe 2023 smoke episodes (Noah et al., 2023; Chen et al., 2025). These risks are not distributed uniformly. Exposure reflects emissions, atmospheric transport, boundary-layer dynamics, and the physical landscape in which people live."
    )
    add_body(
        document,
        "Most public air-quality interfaces compress this complexity into a colored surface or a monitor value. That reduction is useful for communication, but it can imply more spatial certainty than the underlying observation network supports. A city label may hide a valley, ridge, or plateau. A modeled cell may be treated as a direct measurement. Smoke visible aloft may or may not correspond to surface PM2.5. A strong-looking pattern may disappear when tested on a future time block. These are not cosmetic caveats; they determine whether a map is evidence, hypothesis generation, or overclaiming."
    )
    add_body(
        document,
        "The original PASS Equity Atlas prototype addressed provenance, monitor confidence, health burden, and equity context, but its research argument centered on implementation breadth. That framing made it difficult to identify one new scientific contribution. This revised work narrows the claim. The platform is restructured around a specific question: when a user selects a place, does topography explain a reproducible local PM2.5 contrast beyond smoke and meteorology, and can that test be exposed transparently in a public interface?"
    )
    add_body(
        document,
        "The answer is allowed to be no. This is the defining design choice. PASS Terrain-Aware Smoke Atlas separates an interpretable descriptive metric from a predictive test, reports source and cache status, and preserves negative model lift. The resulting platform is not presented as a diagnostic tool, a causal model, or a replacement for regulatory monitoring. It is a research instrument for generating and falsifying terrain-smoke hypotheses."
    )

    add_heading(document, "2. Literature Review", 1)
    add_heading(document, "2.1 Wildfire smoke, PM2.5, and asthma", 2)
    add_body(
        document,
        "Wildfire smoke contains a complex mixture of gases and particles, with PM2.5 frequently used as the principal exposure indicator in epidemiologic work. Noah et al. (2023) summarized consistent acute respiratory effects, particularly among people with asthma, while noting that exposure estimates commonly combine monitors, satellite observations, meteorology, and statistical models. The literature therefore already supports the health relevance of smoke-related PM2.5. The unresolved problem for this project is spatial interpretation: a regional smoke plume does not specify which local topographic positions experience the highest surface concentration."
    )
    add_body(
        document,
        "Chen et al. (2025) demonstrated that extreme smoke events can produce measurable population-level health burdens on short time scales. Such findings motivate rapid and accessible situational awareness, but they do not justify converting a modeled neighborhood surface into individual medical advice. PASS uses asthma sensitivity to frame cautious exposure communication, while explicitly labeling the platform as contextual environmental information rather than diagnosis or treatment guidance."
    )

    add_heading(document, "2.2 Complex terrain, inversion, and local pollution contrasts", 2)
    add_body(
        document,
        "Complex terrain can structure air pollution through cold-air pooling, slope flows, channeling, and inversion dynamics. In the Pittsburgh region, Shmool et al. (2014) used saturation sampling across an inversion-prone urban landscape and found that several pollutants were elevated at lower-elevation locations. A related PM2.5 constituent study showed that terrain-aware sampling can reveal spatial patterns that a sparse regulatory network may miss. These studies establish that valley-ridge differences are plausible and measurable. They also caution against assuming that the same contrast will appear at every place or time."
    )
    add_body(
        document,
        "Terrain does not act alone. Wind direction, mixing height, emission source location, humidity, synoptic conditions, and plume injection height can reverse or overwhelm a simple lowland-versus-highland pattern. For wildfire smoke, elevated transport may pass over a valley without mixing to the surface, while stable conditions may trap locally produced particles. Consequently, elevation is better treated as a context-dependent feature than as a universal exposure correction."
    )

    add_heading(document, "2.3 Terrain-aware PM2.5 estimation", 2)
    add_body(
        document,
        "Recent exposure-modeling research has incorporated elevation, boundary-layer variables, fire activity, satellite aerosol products, and land-use information. Reid et al. (2021) produced daily PM2.5 estimates for western U.S. counties, ZIP codes, and census tracts and showed that performance depends strongly on the validation strategy. Swanson et al. (2022) produced terrain-resolving 1-km PM2.5 maps for the western United States and explicitly represented inversion potential and topographic structure. These efforts demonstrate that terrain-aware prediction is technically feasible at scale."
    )
    add_body(
        document,
        "PASS does not compete with those mature exposure products. Instead, it translates a related scientific principle into an inspectable local experiment. The user can see the terrain cells, PM2.5 samples, smoke polygons, descriptive contrast, and ablation result in one interface. The contribution is therefore methodological and translational: it connects geospatial evidence assembly to a falsification-oriented interaction design."
    )

    add_heading(document, "2.4 Machine learning, validation, and leakage", 2)
    add_body(
        document,
        "Random train-test splits often exaggerate performance for spatiotemporal environmental data because nearby samples share location, weather, and temporal structure. Roberts et al. (2017) formalized the need for validation that reflects the intended prediction task. Just et al. (2020) likewise emphasized evaluation strategies for large-region satellite PM2.5 models, and Boser (2024) showed how split design can reverse conclusions in spatiotemporal machine learning. A model that predicts held-out rows from the same hour is not equivalent to a model that predicts a future period."
    )
    add_body(
        document,
        "The PASS ablation therefore withholds the final 24 UTC hours as an intact temporal block. This is still an early validation design: all grid cells remain near the selected city, and the system has not yet been tested under leave-one-city-out validation. Nevertheless, it directly addresses the most immediate leakage risk and makes the validation boundary visible to the user."
    )

    add_heading(document, "2.5 Research gap", 2)
    add_body(
        document,
        "The reviewed literature supports four observations: wildfire PM2.5 affects respiratory health; complex terrain can structure local pollution; topography can improve exposure models; and validation design materially changes model claims. What remains uncommon is a public-facing system that combines these observations into a transparent local hypothesis test. Typical dashboards show conditions but not an ablation. Typical models report performance but do not let a user inspect the local cells and provenance. Typical consumer agents summarize outputs but may obscure whether values are measured, modeled, cached, or synthetic."
    )
    add_body(
        document,
        "PASS addresses this interface-method gap. The novelty claim is deliberately scoped: to our knowledge from this targeted literature review, the combined workflow of an interpretable lowland-highland exposure differential, smoke-overhead stratification, temporal holdout ablation, and source-status visualization has not been operationalized as a single public environmental-health workbench. A systematic review would be required before making a stronger priority claim."
    )

    add_heading(document, "3. Research Contribution and Questions", 1)
    add_body(
        document,
        "The revised project contributes a falsifiable terrain-smoke analysis module while preserving the broader PASS functions for air quality, health vulnerability, comparison, clinic communication, watchlists, alerts, and sensor planning. Its research value comes from a narrow analytic core rather than from the number of interface tabs."
    )
    for item in (
        "Orographic Exposure Differential (OED): an interpretable local statistic that compares median PM2.5 in lowland and highland thirds of a sampled terrain grid.",
        "Smoke-aware stratification: hourly OED is interpreted alongside NOAA HMS smoke polygons, retaining the distinction between overhead smoke and verified surface smoke.",
        "Terrain ablation: a baseline random forest is compared with an otherwise matched terrain-augmented model on an intact future time block.",
        "Uncertainty-visible interaction: source, timestamp, cache/fallback status, limitations, sample size, validation split, and negative lift are shown in the same research surface.",
        "Agent grounding: the Exposure Navigator can summarize tool outputs, but computed evidence remains visible and provider failures do not alter the underlying measurements.",
    ):
        add_list_item(document, item, bullet_num_id)

    add_heading(document, "3.1 Research questions", 2)
    questions = (
        "RQ1. Within a selected study area, what is the current association between elevation and PM2.5, and how large is the lowland-highland OED?",
        "RQ2. How does hourly OED vary during periods with and without NOAA-observed smoke overhead?",
        "RQ3. Do elevation, topographic position, and ruggedness improve PM2.5 prediction beyond meteorology, smoke indicators, and time-of-day features on a held-out future block?",
        "RQ4. Can users correctly distinguish measured, modeled, cached, fallback, and unavailable evidence when interpreting a local result?",
    )
    for item in questions:
        add_list_item(document, item, decimal_num_id)

    add_heading(document, "3.2 Hypotheses and null-preserving logic", 2)
    add_body(
        document,
        "H1 predicts a positive OED under stable conditions, meaning higher PM2.5 in lower terrain. H2 predicts that the magnitude and sign of OED will differ between smoke-overhead and non-smoke periods. H3 predicts lower held-out RMSE for the terrain-augmented model. These hypotheses are not hard-coded conclusions. A negative OED, weak correlation, or negative predictive lift is retained as a valid result and becomes part of the research record."
    )

    add_heading(document, "4. Methods", 1)
    add_heading(document, "4.1 Study-area construction", 2)
    add_body(
        document,
        "A query begins with a user-selected coordinate and radius of 20, 40, or 60 kilometers. The engine constructs a circular sample from a seven-by-seven lattice and retains 29 cells whose centers fall inside the circular mask. This design provides a stable, legible set of local observations without implying parcel-scale resolution. The grid is generated deterministically, so the same coordinate and radius produce the same cell geometry."
    )
    add_body(
        document,
        "Elevation is requested from the Open-Meteo Elevation API, which derives values from a 90-meter digital elevation model. Cells are ordered by elevation and divided into lower, middle, and upper thirds. These groups are labeled lowland, transition, and highland. The labels are relative to the selected study area, not absolute physiographic classifications. A 300-meter point can therefore be a local highland in one query and a lowland in another."
    )

    add_heading(document, "4.2 Environmental inputs", 2)
    add_table(
        document,
        ["Input", "Operational source", "Role", "Interpretive boundary"],
        [
            ("Elevation", "Open-Meteo Elevation API", "Cell elevation, relief, terrain class", "Modeled DEM, not a field survey"),
            ("PM2.5 and U.S. AQI", "Open-Meteo Air Quality API / CAMS", "Current and historical cell outcomes", "Modeled atmospheric product; not a regulatory monitor"),
            ("Wind and boundary layer", "Open-Meteo Weather API", "Baseline meteorologic predictors", "Grid-model variables"),
            ("Smoke polygons", "NOAA Hazard Mapping System", "Light/medium/heavy overhead context", "Visible smoke extent; not proof of surface concentration"),
            ("Map terrain", "MapLibre raster DEM", "Interactive 3D context and hillshade", "Visualization context, separate from analytic elevation requests"),
        ],
        [1500, 2100, 2400, 3360],
    )
    add_caption(document, "Table 1. Data inputs and the claims they do and do not support.")

    add_heading(document, "4.3 Orographic Exposure Differential", 2)
    add_body(
        document,
        "For hour t, let L be the set of lowland cells and H the set of highland cells. The Orographic Exposure Differential is defined as OED(t) = median[PM2.5(i,t), i in L] - median[PM2.5(j,t), j in H]. Positive OED indicates higher modeled PM2.5 in the local lowland third; negative OED indicates higher modeled PM2.5 in the highland third. Median aggregation reduces sensitivity to a single extreme cell."
    )
    add_body(
        document,
        "OED is descriptive and local. It does not estimate a causal terrain effect because lowland and highland cells may also differ in roads, population density, emissions, and land use. The interface therefore pairs OED with the Spearman rank association between cell elevation and PM2.5, relief, sample count, and source status. Agreement between OED and the rank coefficient strengthens descriptive coherence but does not remove confounding."
    )

    add_heading(document, "4.4 Smoke-overhead classification", 2)
    add_body(
        document,
        "The engine retrieves daily NOAA HMS KML polygons and tests whether the study center or sampled cells intersect light, medium, or heavy smoke geometry. The highest intersecting density is reported for each day. Because HMS identifies smoke observed in satellite imagery, the system uses the phrase smoke overhead rather than surface smoke exposure. The PM2.5 field remains the surface-concentration indicator. Days without a retrievable polygon are marked unavailable rather than smoke-free."
    )

    add_heading(document, "4.5 Random-forest ablation", 2)
    add_body(
        document,
        "The baseline feature set contains wind speed, planetary boundary-layer height, smoke density or availability, and cyclic hour-of-day terms. The terrain model adds elevation, topographic position index, and ruggedness. Both models use the same sampled rows and random-forest implementation. Training uses every sixth hour from the pre-holdout period to reduce serial redundancy and computation; all observations from the final 24 UTC hours are withheld for testing."
    )
    add_body(
        document,
        "Performance is reported with root mean squared error (RMSE), coefficient of determination (R-squared), and terrain lift: 100 x (RMSE_baseline - RMSE_terrain) / RMSE_baseline. Positive lift means the terrain model reduced error; zero or negative lift means the added terrain features did not improve this holdout. Permutation importance is calculated by shuffling each terrain feature in the held-out rows and measuring the change in RMSE."
    )
    add_body(
        document,
        "This ablation tests incremental prediction, not causality. Random forests can capture nonlinear interactions but do not establish why a feature predicts. Feature importance can also be unstable when predictors are correlated. The interface therefore treats the ablation as one line of evidence and exposes the baseline, test-window definition, sample counts, and failure states."
    )

    add_heading(document, "4.6 Caching, provenance, and failure handling", 2)
    add_body(
        document,
        "Terrain-smoke analyses are cached for 30 minutes to limit repeated upstream calls. Each response includes a source ledger with provider, variable, timestamp, status, and limitation. Partial failure is represented explicitly: for example, PM2.5 may remain available when a NOAA file is missing. The API does not silently convert missing smoke polygons into a no-smoke observation. Provider errors are surfaced in the evidence record while the rest of the page remains usable."
    )

    add_heading(document, "4.7 Ethical and clinical boundary", 2)
    add_body(
        document,
        "The application is intended for research exploration and situational context. It is not a medical device, exposure diagnosis, evacuation system, or substitute for AirNow, local public-health guidance, emergency services, or a clinician. Health-related language is conservative and should be evaluated with community partners, asthma clinicians, and risk-communication experts before deployment in a care setting."
    )

    add_heading(document, "5. System Design and Implementation", 1)
    add_body(
        document,
        "PASS is implemented as a Next.js application with typed server routes, a MapLibre client map, deterministic analytic functions, and an optional large-language-model navigator. The terrain-smoke engine executes on the server so provider credentials and data integration remain outside the browser. The public endpoint accepts latitude, longitude, study radius, and historical window parameters, validates bounds, and returns structured JSON."
    )
    add_body(
        document,
        "The Terrain & Smoke Lab preserves the operational style of the existing platform. Users can search a location, switch between kilometers and miles, choose a 20/40/60-kilometer study radius, and select a three-, five-, or seven-day window. Map controls toggle terrain classes, PM2.5 samples, hillshade, and NOAA plume geometry. The evidence panel reports current OED, historical OED distribution, rank correlation, random-forest ablation, smoke context, and source records."
    )
    document.add_picture(str(SCREENSHOT_PATH), width=Inches(6.5))
    add_caption(
        document,
        "Figure 1. PASS Terrain & Smoke Lab. The interface combines 3D terrain context, cell-level PM2.5, OED, temporal validation, smoke status, and an evidence ledger. Screenshot captured from the implemented August 2026 prototype."
    )
    add_body(
        document,
        "The Exposure Navigator is intentionally downstream of the analytic engine. Its role is to explain computed evidence and help users formulate comparisons or sensor-planning questions. A provider cascade can use Groq or Gemini when configured, with a deterministic router as fallback. The badge identifies the response mode, and tool calls remain inspectable. The language model cannot replace failed upstream measurements or modify the OED and model metrics."
    )

    add_heading(document, "6. Worked Example: Allentown, Pennsylvania", 1)
    add_body(
        document,
        "A worked query centered on Allentown used a 40-kilometer study radius and a three-day history. The local landscape was classified as mountainous, with 331 meters of relief across the 29-cell grid. The current lowland median PM2.5 was 8.8 micrograms per cubic meter, compared with 3.6 in the highland group, producing an OED of +5.2. The Spearman elevation-PM2.5 coefficient was -0.43, consistent with higher current modeled concentrations at lower elevations."
    )
    add_body(
        document,
        "The historical picture was more modest. Median OED across the window was +0.4 micrograms per cubic meter, and 67% of hourly contrasts were positive. NOAA HMS indicated light smoke overhead on the latest available day and smoke intersection on all three retrieved days. These observations generate a plausible hypothesis that terrain and mixing conditions may have structured the current surface, but they do not isolate smoke from local sources or meteorology."
    )
    add_body(
        document,
        "The predictive test did not support a generalized terrain benefit for this window. Baseline held-out RMSE was 7.93, while terrain-augmented RMSE was 8.58. Terrain lift was therefore -8.2%. The training set contained 232 sampled rows and the final 24-hour test block contained 696 rows. This negative result is scientifically important: a visible current lowland-highland contrast can coexist with poor future predictive utility. The interface states both findings instead of presenting only the visually compelling one."
    )
    document.add_picture(str(RESULTS_FIGURE_PATH), width=Inches(6.5))
    add_caption(
        document,
        "Figure 2. Worked Allentown example. The current OED and held-out ablation answer different questions. Values are a time-specific prototype response and should be regenerated for any subsequent analysis."
    )
    add_table(
        document,
        ["Evidence", "Result", "Interpretation"],
        [
            ("Landscape relief", "331 m", "Large enough to motivate a local terrain test"),
            ("Current OED", "+5.2 ug/m3", "Lowland median exceeded highland median at the current snapshot"),
            ("Spearman rho", "-0.43", "Moderate inverse elevation-PM2.5 rank association"),
            ("Historical median OED", "+0.4 ug/m3", "Typical contrast was much smaller than the snapshot"),
            ("Positive OED hours", "67%", "Lowland excess was common, not universal"),
            ("Terrain lift", "-8.2%", "Terrain features worsened final-block prediction"),
            ("Smoke context", "Light; 3/3 days", "Overhead smoke was observed, not proven at the surface"),
        ],
        [2400, 1800, 5160],
        numeric_columns={1},
    )
    add_caption(document, "Table 2. Interpretation of the Allentown worked example.")

    add_heading(document, "7. Discussion", 1)
    add_heading(document, "7.1 What is new", 2)
    add_body(
        document,
        "The project does not claim novelty for mapping AQI, using terrain in a PM2.5 model, applying random forests, or showing wildfire smoke polygons. Each has precedent. The new contribution is their arrangement into an uncertainty-first local research protocol. OED makes a terrain contrast interpretable; smoke stratification provides event context; ablation asks whether terrain improves prediction; and the interface exposes enough provenance to challenge the result."
    )
    add_body(
        document,
        "This arrangement changes the app from a feature inventory into a research instrument. A user can move from observation to test: identify a valley-ridge contrast, inspect whether smoke was present, compare multiple hours, evaluate predictive lift, and export a source trail. The workflow can support hypothesis generation for temporary sensor placement, especially where regulatory monitors are sparse, but it does not claim to optimize placement without field validation."
    )

    add_heading(document, "7.2 Why the negative result matters", 2)
    add_body(
        document,
        "Environmental interfaces are vulnerable to visual confirmation bias. Color gradients, 3D terrain, and plume outlines can make a pattern feel more certain than it is. The Allentown example demonstrates the opposite design principle. The current map suggests a lowland excess, yet the future-block ablation finds no predictive improvement. Showing this disagreement teaches users that descriptive association and generalizable prediction are distinct."
    )
    add_body(
        document,
        "Negative lift may reflect short history, coarse model inputs, spatial confounding, limited feature engineering, time-varying atmospheric processes, or genuine absence of a stable terrain contribution. The correct response is not to hide the result or add model complexity until lift becomes positive. It is to expand the evaluation across places, seasons, smoke regimes, and monitor observations."
    )

    add_heading(document, "7.3 Relevance to asthma-sensitive communities", 2)
    add_body(
        document,
        "Asthma-sensitive residents and public-health teams need timely information, but they also need calibrated uncertainty. A terrain-aware view may reveal where a citywide average is least representative and where temporary monitoring could reduce uncertainty. Equity layers can prioritize communities with high health burden and weak monitor coverage. However, vulnerability scores should not be used to infer an individual's exposure or clinical status. Community governance is required when translating these metrics into outreach or resource allocation."
    )

    add_heading(document, "7.4 Research versus operations", 2)
    add_body(
        document,
        "PASS now contains two complementary modes. Operational views summarize current conditions, forecasts, watchlists, alerts, clinic-safe notes, and comparisons. The Terrain & Smoke Lab is explicitly analytical: it slows the user down, identifies assumptions, and reports tests. Keeping both modes in one platform allows research outputs to inform operations while maintaining a visible boundary between exploratory evidence and action guidance."
    )

    add_heading(document, "8. Limitations", 1)
    limitations = (
        "PM2.5 and AQI values are modeled atmospheric fields, not direct regulatory monitor observations. Apparent cell-level precision may exceed real uncertainty.",
        "NOAA HMS polygons indicate satellite-observed smoke extent and density category. They do not prove smoke reached the surface or quantify source-specific PM2.5.",
        "The 29-cell relative terrain classification is query-dependent and does not encode full watershed, valley orientation, land cover, road emissions, or building morphology.",
        "OED is descriptive and susceptible to spatial confounding. It should not be interpreted as a causal effect of elevation.",
        "The current holdout is temporal within one local grid. It addresses hour-level leakage but does not establish transfer to a new city or season.",
        "Random-forest hyperparameters are operational defaults rather than tuned through nested cross-validation. Permutation importance may be unstable under correlated features.",
        "The worked example is time-specific and was generated during development. It is evidence that the workflow runs, not a population-level result.",
        "The user-comprehension question has not yet been evaluated. Provenance may be visible without being understood.",
        "The system has no completed clinical, accessibility, community-governance, or emergency-response validation and must not be used for medical decisions.",
    )
    for item in limitations:
        add_list_item(document, item, bullet_num_id)

    add_heading(document, "9. Validation Roadmap", 1)
    add_body(
        document,
        "The next study should preregister locations, time windows, primary metrics, and exclusion rules before looking at results. Candidate sites should include inversion-prone valleys, flatter comparison cities, smoke-prone western communities, and eastern urban corridors. A minimum design would span four seasons and multiple smoke regimes."
    )
    roadmap = (
        "Monitor anchoring. Join EPA Air Quality System or AirNow observations and evaluate calibration, bias, and OED agreement at monitored sites.",
        "Blocked validation. Use leave-one-day, leave-one-event, leave-one-city, and spatial-block tests; compare them with random splits to quantify optimism.",
        "Model comparison. Benchmark the random forest against regularized linear models, gradient boosting, and a no-terrain persistence model. Keep terrain only when it improves prespecified metrics.",
        "Smoke attribution. Add plume transport, fire radiative power, aerosol optical depth, and wind-direction features while retaining the distinction between smoke aloft and surface concentration.",
        "Sensitivity analysis. Vary study radius, grid density, terrain quantile thresholds, elevation source, holdout length, and missing-polygon rules.",
        "Field pilot. Deploy temporary sensors across paired lowland and highland sites selected before an event and compare observed versus predicted contrasts.",
        "Human-factors evaluation. Test whether residents, clinicians, and public-health staff correctly identify modeled versus measured data, uncertainty, and negative lift.",
        "Equity governance. Co-design prioritization criteria and audit whether the interface benefits communities with high burden and low monitor coverage without stigmatizing them.",
    )
    roadmap_num_id = add_numbering_definition(document, bullet=False)
    for item in roadmap:
        add_list_item(document, item, roadmap_num_id)

    add_heading(document, "9.1 Proposed primary endpoints", 2)
    add_table(
        document,
        ["Question", "Primary endpoint", "Validation design", "Decision rule"],
        [
            ("Terrain contrast", "Monitor-anchored OED error", "Paired lowland/highland sensors", "Report magnitude and interval; no sign-only success criterion"),
            ("Predictive value", "Delta RMSE versus baseline", "Leave-one-city/event-out", "Terrain retained only if improvement is positive and stable"),
            ("Smoke context", "Surface PM2.5 discrimination", "HMS-stratified monitor hours", "Separate unavailable polygons from no-smoke days"),
            ("User understanding", "Correct provenance interpretation", "Task-based usability study", "Predefined comprehension and confidence calibration targets"),
        ],
        [1600, 2100, 2700, 2960],
    )
    add_caption(document, "Table 3. Proposed endpoints for a preregistered validation study.")

    add_heading(document, "10. Conclusion", 1)
    add_body(
        document,
        "PASS Terrain-Aware Smoke Atlas reframes an environmental-health website around one testable contribution: an accessible, source-visible method for asking whether local terrain structures PM2.5 during smoke-relevant periods. The platform computes an interpretable OED, displays cell-level terrain and pollution context, and requires terrain features to demonstrate predictive lift on an untouched future block."
    )
    add_body(
        document,
        "The worked Allentown example illustrates the intended scientific behavior. A strong current terrain contrast appeared, while the terrain model performed worse than the baseline on held-out data. Retaining both results is more valuable than presenting a seamless but overstated story. The next contribution must come from preregistered multi-city, monitor-anchored validation and community-centered evaluation. Until then, PASS is best understood as a transparent hypothesis-generation workbench, not a validated exposure or clinical decision system."
    )

    add_heading(document, "Data and Code Availability", 1)
    add_body(
        document,
        "The application source is maintained at https://github.com/noyo12394/asthamaaa. The public deployment is available at https://asthamaaa.vercel.app/. The repository documents source endpoints, cache behavior, environment variables, and the terrain-smoke API. External provider terms and rate limits apply."
    )

    add_heading(document, "Acknowledgments", 1)
    add_body(
        document,
        "This draft was developed as part of the PASS project at Lehigh University. The author thanks reviewers who emphasized the need to distinguish implementation breadth from a clear research contribution."
    )

    add_heading(document, "References", 1)
    references = [
        "Boser, A. S. (2024). Validating spatio-temporal environmental machine learning models: Simpson's paradox and data splits. Environmental Research Communications, 6(3), 031003. https://doi.org/10.1088/2515-7620/ad2e44",
        "Chen, H., Kaufman, J. S., Chen, C., Wang, J., Maier, A., van Dijk, A., Slipp, N., Rana, J., MacIntyre, E., Su, Y. S., Kim, J. H., & Benmarhnia, T. (2025). Impact of the 2023 wildfire smoke episodes in Ontario, Canada, on asthma and other health outcomes: An interrupted time-series analysis. CMAJ, 197, E465-E477. https://doi.org/10.1503/cmaj.241506",
        "Just, A. C., Arfer, K. B., Rush, J., Dorman, M., Shtein, A., Lyapustin, A., & Kloog, I. (2020). Advancing methodologies for applying machine learning and evaluating spatiotemporal models of fine particulate matter (PM2.5) using satellite data over large regions. Atmospheric Environment, 239, 117649. https://doi.org/10.1016/j.atmosenv.2020.117649",
        "NOAA Office of Satellite and Product Operations. (2026). Hazard Mapping System fire and smoke product. https://www.ospo.noaa.gov/products/land/hms.html",
        "Noah, T. L., Worden, C. P., Rebuli, M. E., & Jaspers, I. (2023). The effects of wildfire smoke on asthma and allergy. Current Allergy and Asthma Reports. https://doi.org/10.1007/s11882-023-01090-1",
        "Open-Meteo. (2026a). Air Quality API documentation. https://open-meteo.com/en/docs/air-quality-api",
        "Open-Meteo. (2026b). Elevation API documentation. https://open-meteo.com/en/docs/elevation-api",
        "Open-Meteo. (2026c). Weather Forecast API documentation. https://open-meteo.com/en/docs",
        "Reid, C. E., Considine, E. M., Maestas, M. M., & Li, G. (2021). Daily PM2.5 concentration estimates by county, ZIP code, and census tract in 11 western states 2008-2018. Scientific Data, 8, 112. https://doi.org/10.1038/s41597-021-00891-1",
        "Roberts, D. R., Bahn, V., Ciuti, S., Boyce, M. S., Elith, J., Guillera-Arroita, G., Hauenstein, S., Lahoz-Monfort, J. J., Schroder, B., Thuiller, W., Warton, D. I., Wintle, B. A., Hartig, F., & Dormann, C. F. (2017). Cross-validation strategies for data with temporal, spatial, hierarchical, or phylogenetic structure. Ecography, 40(8), 913-929. https://doi.org/10.1111/ecog.02881",
        "Shmool, J. L. C., Michanowicz, D. R., Cambal, L., Tunno, B., Howell, J., Gillooly, S., Roper, C., Tripathy, S., Chubb, L. G., Eisl, H. M., & Clougherty, J. E. (2014). Saturation sampling for spatial variation in multiple air pollutants across an inversion-prone metropolitan area of complex terrain. Environmental Health, 13, 28. https://doi.org/10.1186/1476-069X-13-28",
        "Swanson, A., Holden, Z. A., Graham, J., Warren, D. A., Noonan, C., & Landguth, E. (2022). Daily 1 km terrain resolving maps of surface fine particulate matter for the western United States 2003-2021. Scientific Data, 9, 466. https://doi.org/10.1038/s41597-022-01488-y",
        "U.S. Geological Survey. (2026). Elevation Point Query Service accuracy guidance. https://www.usgs.gov/faqs/how-accurate-are-elevations-generated-elevation-point-query-service-national-map",
    ]
    for reference in references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.first_line_indent = Inches(-0.35)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(reference)
        run.font.name = "Arial"
        run.font.size = Pt(10)

    properties = document.core_properties
    properties.title = "PASS Terrain-Aware Smoke Atlas"
    properties.subject = "Terrain-aware PM2.5, wildfire smoke, uncertainty, and public environmental-health visualization"
    properties.author = "Noyonica Chatterjee"
    properties.keywords = "PM2.5, wildfire smoke, terrain, asthma, random forest, geospatial visualization"
    properties.comments = "Research manuscript draft generated for the PASS project."

    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
